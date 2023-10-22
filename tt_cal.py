import docx
from ics import Calendar, Event
from datetime import datetime, timedelta
import re

def splittimerange(time_range):
    match = re.search(r'(\d+:\d+[ap]m) – (\d+:\d+[ap]m)', time_range)
    if match:
        start_time, end_time = match.groups()
        return start_time, end_time
    else:
        return None, None
    
def splittime(time_str):
    return datetime.strptime(time_str, '%I:%M%p').time()

def create_event(day, time_range, summary):
    start_time, end_time = splittimerange(time_range)
    if start_time is not None and end_time is not None:
        event = Event()
        event.name = summary
        event.start = start_date + splittime(start_time)
        event.end = start_date + splittime(end_time)
        event.rrule = f"FREQ=WEEKLY;BYDAY={day.upper()};UNTIL={end_date.strftime('%Y%m%d')}"
        event.reminders = {'useDefault': False}
        return event

def main(docx_path, start_date, end_date, ics_filename):
    cal = Calendar()

    doc = docx.Document(docx_path)

    start_date = datetime.strptime(start_date, '%Y-%m-%d')
    end_date = datetime.strptime(end_date, '%Y-%m-%d')

    table = doc.tables[0]
    days = [cell.text.strip() for cell in table.cell(0, 1).paragraphs]
    times = [cell.text.strip() for cell in table.column_cells(0)[1:]]
    for i, row in enumerate(table.rows[1:]):
        day = row.cells[0].text
        for j, cell in enumerate(row.cells[1:]):
            summary = cell.text
            if summary:
                event_date = start_date + timedelta(days=i)
                if event_date <= end_date:
                    event = create_event(day, times[j], summary)
                    if event:
                        cal.events.add(event)

    with open(ics_filename, 'w') as ics_file:
        ics_file.writelines(cal)

if __name__ == "__main__":
    docx_path = input("Enter the .docx path: ")
    docx_path = docx_path.replace("\\", "/")
    start_date = input("Enter the Start Date: [YYYY-MM-DD]: ")
    end_date = input("Enter the End Date: [YYYY-MM-DD]: ")
    ics_filename = "output.ics"

    main(docx_path, start_date, end_date, ics_filename)